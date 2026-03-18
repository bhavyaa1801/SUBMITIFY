from dotenv import load_dotenv
load_dotenv()
from fastapi import FastAPI , Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel , Field
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import List
import pdfkit
import os
import html
import re
import base64
from io import BytesIO
from jinja2 import Environment, FileSystemLoader
from experiment_generator import generate_experiments
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded
import asyncio
from datetime import datetime, timedelta

app = FastAPI()

origins = os.getenv("ALLOWED_ORIGINS", "http://localhost:5173").split(",")

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

os.makedirs("generated", exist_ok=True)
app.mount("/files", StaticFiles(directory="generated"), name="files")


limiter = Limiter(key_func=get_remote_address)
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

# ── Request Models ──────────────────────────────────────────

class ProjectInfo(BaseModel):
    university:   str = Field(..., max_length=250)
    year:         str = Field(..., max_length=20)
    department:   str = Field(..., max_length=200)
    subject:      str = Field(..., max_length=200)
    code:         str = Field(..., max_length=50)
    student:      str = Field(..., max_length=150)
    roll:         str = Field(..., max_length=50)
    course:       str = Field(..., max_length=100)
    semester:     str = Field(..., max_length=20)
    submitted_to: str = Field(..., max_length=150)
    designation:  str = Field(..., max_length=100)
    logo:         str | None = None

class GenerateRequest(BaseModel):
    project:   ProjectInfo
    language:  str = Field(..., max_length=50)
    questions: List[str] = Field(..., max_length=25)


# ── Helpers ─────────────────────────────────────────────────

async def cleanup_generated_folder():
    """Delete files older than 1 hour from generated/ every 30 minutes."""
    while True:
        await asyncio.sleep(1800)  
        try:
            now = datetime.now()
            folder = "generated"
            deleted = 0
            for filename in os.listdir(folder):
                if filename == "sample.pdf":
                    continue
                filepath = os.path.join(folder, filename)
                if os.path.isfile(filepath):
                    file_age = now - datetime.fromtimestamp(os.path.getmtime(filepath))
                    if file_age > timedelta(hours=1):
                        os.remove(filepath)
                        deleted += 1
                        print(f"Deleted old file: {filename}")
            if deleted:
                print(f"Cleanup done — {deleted} files deleted")
            else:
                print("Cleanup done — nothing to delete")
        except Exception as e:
            print(f"Cleanup error: {e}")

def set_cell_bg(cell, hex_color: str):
    """Set background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)

def bold_cell(cell, text, font_size=11, color=None):
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def add_horizontal_line(doc):
    """Add a thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)


def build_docx(project, experiments, docx_path):
    doc = DocxDocument()

    # ── Page margins ──────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Default font ──────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ════════════════════════════════════════════════════
    # COVER PAGE
    # ════════════════════════════════════════════════════



    # University name
    univ_para = doc.add_paragraph()
    univ_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    univ_run = univ_para.add_run(project.university.upper())
    univ_run.bold = True
    univ_run.font.size = Pt(16)
    univ_run.font.color.rgb = RGBColor(0, 0, 0)
    univ_para.paragraph_format.space_after = Pt(2)

    # Logo
    if project.logo:
        try:
            header_data = project.logo.split(",")[1] if "," in project.logo else project.logo
            img_bytes = base64.b64decode(header_data)
            img_stream = BytesIO(img_bytes)
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = logo_para.add_run()
            run.add_picture(img_stream, width=Inches(1.2))
            logo_para.paragraph_format.space_after = Pt(6)
        except Exception as e:
            print("Logo error:", e)

    # Department + Year
    for text in [f"Department of {project.department}", f"Academic Year: {project.year}"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.space_after = Pt(2)

    # Spacer
    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # Practical File title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("Practical File")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    title_para.paragraph_format.space_after = Pt(4)

    # Subject name
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(project.subject)
    sub_run.bold = True
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = RGBColor(0, 0, 0)
    sub_para.paragraph_format.space_after = Pt(2)

    # Subject code
    code_para = doc.add_paragraph()
    code_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    code_run = code_para.add_run(f"(Subject Code: {project.code})")
    code_run.font.size = Pt(11)
    code_run.font.color.rgb = RGBColor(0, 0, 0)
    code_para.paragraph_format.space_after = Pt(32)

    # ── Submitted to / by — invisible two-column table ──
    info_table = doc.add_table(rows=1, cols=2)

    # Remove ALL borders from the table itself
    tbl = info_table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        tblBorders.append(border)
    tblPr.append(tblBorders)

    left_cell  = info_table.rows[0].cells[0]
    right_cell = info_table.rows[0].cells[1]

    # Also remove cell borders
    for cell in [left_cell, right_cell]:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for side in ["top", "left", "bottom", "right"]:
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "none")
            b.set(qn("w:sz"), "0")
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "auto")
            tcBorders.append(b)
        tcPr.append(tcBorders)
        set_cell_margins(cell, top=80, bottom=80, left=0, right=0)

    def add_cover_line(cell, label, value, align=WD_ALIGN_PARAGRAPH.LEFT):
        p = cell.add_paragraph()
        p.alignment = align
        label_run = p.add_run(f"{label} : ")
        label_run.bold = True
        label_run.font.size = Pt(11)
        label_run.font.color.rgb = RGBColor(0, 0, 0)
        val_run = p.add_run(value)
        val_run.font.size = Pt(11)
        val_run.font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.space_after = Pt(3)

    left_cell.paragraphs[0].clear()
    add_cover_line(left_cell, "Submitted to", project.submitted_to)
    add_cover_line(left_cell, "Designation",  project.designation)
    add_cover_line(left_cell, "Department",   project.department)

    right_cell.paragraphs[0].clear()
    add_cover_line(right_cell, "Submitted by",   project.student,  WD_ALIGN_PARAGRAPH.RIGHT)
    add_cover_line(right_cell, "Enrollment No",  project.roll,     WD_ALIGN_PARAGRAPH.RIGHT)
    add_cover_line(right_cell, "Course",         project.course,   WD_ALIGN_PARAGRAPH.RIGHT)
    add_cover_line(right_cell, "Semester",       project.semester, WD_ALIGN_PARAGRAPH.RIGHT)

    doc.add_page_break()

    # ════════════════════════════════════════════════════
    # INDEX PAGE
    # ════════════════════════════════════════════════════

    idx_heading = doc.add_paragraph()
    idx_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    idx_run = idx_heading.add_run("Index")
    idx_run.bold = True
    idx_run.font.size = Pt(14)
    idx_run.font.color.rgb = RGBColor(0, 0, 0)
    idx_heading.paragraph_format.space_after = Pt(10)

    idx_table = doc.add_table(rows=1, cols=4)
    idx_table.style = "Table Grid"

    # Header row — plain bold black, no colour fill
    headers = ["E. No.", "Aim", "Date", "Sign"]
    col_widths = [700, 5200, 1300, 1300]
    for i, (cell, hdr) in enumerate(zip(idx_table.rows[0].cells, headers)):
        # Set column width
        tc_w = OxmlElement("w:tcW")
        tc_w.set(qn("w:w"), str(col_widths[i]))
        tc_w.set(qn("w:type"), "dxa")
        cell._tc.get_or_add_tcPr().append(tc_w)

        set_cell_margins(cell)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(hdr)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0, 0, 0)

    # Data rows
    for i, exp in enumerate(experiments, 1):
        row = idx_table.add_row()
        values  = [str(i), exp.get("name", f"Experiment {i}"), "", ""]
        aligns  = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                   WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER]
        for j, (cell, val, align) in enumerate(zip(row.cells, values, aligns)):
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = align
            r = p.add_run(val)
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_page_break()

    # ════════════════════════════════════════════════════
    # EXPERIMENT PAGES
    # ════════════════════════════════════════════════════

    for i, exp in enumerate(experiments, 1):
        name = exp.get("name", f"Experiment {i}")

        # "Experiment N" heading
        exp_para = doc.add_paragraph()
        exp_run = exp_para.add_run(f"Experiment {i}")
        exp_run.bold = True
        exp_run.font.size = Pt(14)
        exp_run.font.color.rgb = RGBColor(0, 0, 0)
        exp_para.paragraph_format.space_after = Pt(2)

        # Experiment name
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(name.title())
        name_run.bold = True
        name_run.font.size = Pt(12)
        name_run.font.color.rgb = RGBColor(0, 0, 0)
        name_para.paragraph_format.space_after = Pt(6)

        sections = [
            ("aim",       "Aim"),
            ("algorithm", "Algorithm"),
            ("code",      "Program"),
            ("output",    "Output"),
            ("result",    "Result"),
        ]

        for key, label in sections:
            content = exp.get(key, "").strip()
            if not content:
                continue

            # Section label
            lbl_para = doc.add_paragraph()
            lbl_run = lbl_para.add_run(label)
            lbl_run.bold = True
            lbl_run.font.size = Pt(11)
            lbl_run.font.color.rgb = RGBColor(0, 0, 0)
            lbl_para.paragraph_format.space_before = Pt(8)
            lbl_para.paragraph_format.space_after = Pt(3)

            if key == "code":
                # Plain monospace, no background
                for line in content.splitlines():
                    code_para = doc.add_paragraph()
                    code_para.paragraph_format.space_before = Pt(0)
                    code_para.paragraph_format.space_after = Pt(0)
                    code_para.paragraph_format.left_indent = Inches(0.2)
                    code_run = code_para.add_run(line if line else " ")
                    code_run.font.name = "Times New Roman"
                    code_run.font.size = Pt(9.5)
                    code_run.font.color.rgb = RGBColor(0, 0, 0)

            elif key == "algorithm":
                for line in content.splitlines():
                    line = line.strip()
                    if not line:
                        continue
                    alg_para = doc.add_paragraph()
                    alg_para.paragraph_format.space_before = Pt(0)
                    alg_para.paragraph_format.space_after = Pt(2)
                    alg_para.paragraph_format.left_indent = Inches(0.2)
                    alg_run = alg_para.add_run(line)
                    alg_run.font.size = Pt(11)
                    alg_run.font.color.rgb = RGBColor(0, 0, 0)

            else:
                content_para = doc.add_paragraph()
                content_para.paragraph_format.space_after = Pt(4)
                content_para.paragraph_format.left_indent = Inches(0.2)
                content_run = content_para.add_run(content)
                content_run.font.size = Pt(11)
                content_run.font.color.rgb = RGBColor(0, 0, 0)
                if key == "output":
                    content_run.font.name = "Times New Roman"
                    content_run.font.size = Pt(10)

        doc.add_page_break()

    doc.save(docx_path)

def sanitize(value: str) -> str:
    """Strip leading/trailing whitespace and escape HTML special characters."""
    if not value:
        return ""
    value = value.strip()
    value = html.escape(value)          # converts & < > " ' to safe entities
    value = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", value)  # remove control characters
    return value

def sanitize_project(project: ProjectInfo) -> ProjectInfo:
    """Return a sanitized copy of project info."""
    return ProjectInfo(
        university   = sanitize(project.university),
        year         = sanitize(project.year),
        department   = sanitize(project.department),
        subject      = sanitize(project.subject),
        code         = sanitize(project.code),
        student      = sanitize(project.student),
        roll         = sanitize(project.roll),
        course       = sanitize(project.course),
        semester     = sanitize(project.semester),
        submitted_to = sanitize(project.submitted_to),
        designation  = sanitize(project.designation),
        logo         = project.logo,  # logo is base64, don't sanitize
    )

# ── Routes ──────────────────────────────────────────────────

@app.on_event("startup")
async def start_cleanup():
    asyncio.create_task(cleanup_generated_folder())
    
@app.post("/api/generate")
@limiter.limit("10/minute")
async def generate(request: Request, req: GenerateRequest):
    req = GenerateRequest(
        project=sanitize_project(req.project),
        language=sanitize(req.language),
        questions=[sanitize(q) for q in req.questions],
    )
    experiments = generate_experiments(
        req.questions,
        req.language,
        req.project.subject
    )
    return {"experiments": experiments}


@app.post("/api/generate-pdf")
@limiter.limit("5/minute")
async def generate_pdf(request: Request, req: GenerateRequest):

    req = GenerateRequest(
        project=sanitize_project(req.project),
        language=sanitize(req.language),
        questions=[sanitize(q) for q in req.questions],
    )

    experiments = generate_experiments(
        req.questions,
        req.language,
        req.project.subject
    )

    # ── Unique filename from user details ─────────────
    def slugify(text):
        text = text.strip().lower()
        text = re.sub(r"[^\w\s-]", "", text)
        text = re.sub(r"\s+", "_", text)
        return text

    enrollment = slugify(req.project.roll)
    name       = slugify(req.project.student)
    subject    = slugify(req.project.subject)

    file_name = f"{enrollment}_{name}_{subject}"
    pdf_path  = f"generated/{file_name}.pdf"
    docx_path = f"generated/{file_name}.docx"

    # ── PDF ───────────────────────────────────────────
    env = Environment(loader=FileSystemLoader("templates"))
    cover_html = env.get_template("cover.html").render(project=req.project.dict())
    index_html = env.get_template("index.html").render(experiments=experiments)
    exp_html   = env.get_template("experiments.html").render(experiments=experiments)

    final_html = f"""
    <html><head><meta charset="UTF-8"></head>
    <body>
      {cover_html}
      <div style="page-break-after: always;"></div>
      {index_html}
      <div style="page-break-after: always;"></div>
      {exp_html}
    </body></html>
    """

    wkhtmltopdf_path = os.getenv("WKHTMLTOPDF_PATH", "/usr/bin/wkhtmltopdf")
    config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_path)
    pdfkit.from_string(
        final_html, pdf_path,       
        configuration=config,
        options={"enable-local-file-access": "", "page-size": "A4"},
        css="styles.css"
    )

    # ── DOCX ──────────────────────────────────────────
    build_docx(req.project, experiments, docx_path)  

    base_url = os.getenv("API_URL", "http://localhost:8000")

    return {
        "pdf_url":  f"{base_url}/files/{file_name}.pdf",
        "docx_url": f"{base_url}/files/{file_name}.docx",
    }